import streamlit as st
import pandas as pd
from docx import Document
import tempfile
import zipfile
import os

# -------------------------------
# Función para generar documentos
# -------------------------------
def generar_documentos(excel_file, word_template):
    df = pd.read_excel(excel_file)

    # Crear carpeta temporal
    temp_dir = tempfile.mkdtemp()

    for idx, row in df.iterrows():
        doc = Document(word_template)

        # Manejo de Subinvestigador
        if pd.notna(row.get("Subinvestigador", "")) and str(row["Subinvestigador"]).strip() != "":
            subinvestigador_texto = "Sub Investigador " + str(row["Subinvestigador"])
        else:
            subinvestigador_texto = ""

        # Diccionario de reemplazos
        reemplazos = {
            "<<NUMERO_PROTOCOLO>>": str(row.get("Numero de protocolo", "")),
            "<<TITULO_ESTUDIO>>": str(row.get("Titulo del Estudio", "")),
            "<<PATROCINADOR>>": str(row.get("Patrocinador", "")),
            "<<INVESTIGADOR>>": str(row.get("Investigador", "")),
            "<<INSTITUCION>>": str(row.get("Institucion", "")),
            "<<DIRECCION>>": str(row.get("Direccion", "")),
            "<<CARGO_INVESTIGADOR>>": str(row.get("Cargo del Investigador en la Institucion", "")),
            "<<TELEFONO_24HS>>": str(row.get("Telefono 24hs", "")),
            "<<COMITE>>": str(row.get("Comite", "")),
            "<<SUBINVESTIGADOR>>": subinvestigador_texto
        }

        # Eliminar párrafos si Dirección contiene "Córdoba"
        if "cordoba" in str(row.get("Direccion", "")).lower():
            parrafos_a_eliminar = [
                "El medico del estudio discutira con Usted que metodo anticonceptivo se considera adecuado.",
                "requerido para centros de la provincia de buenos aires"
            ]
            for p in doc.paragraphs:
                for texto in parrafos_a_eliminar:
                    if p.text.strip().startswith(texto):
                        p.clear()

        # Reemplazar placeholders
        for p in doc.paragraphs:
            for placeholder, valor in reemplazos.items():
                if placeholder in p.text:
                    inline = p.runs
                    for i in range(len(inline)):
                        if placeholder in inline[i].text:
                            inline[i].text = inline[i].text.replace(placeholder, valor)

        # Guardar documento temporal
        nombre_archivo = f"documento_{idx+1}.docx"
        ruta_salida = os.path.join(temp_dir, nombre_archivo)
        doc.save(ruta_salida)

    # Crear un ZIP con todos los documentos
    zip_path = os.path.join(temp_dir, "documentos.zip")
    with zipfile.ZipFile(zip_path, "w") as zipf:
        for archivo in os.listdir(temp_dir):
            if archivo.endswith(".docx"):
                zipf.write(os.path.join(temp_dir, archivo), archivo)

    return zip_path


# -------------------------------
# Interfaz Streamlit
# -------------------------------
st.set_page_config(page_title="Generador de Consentimientos", layout="wide")

st.title("📄 Generador Automático de Consentimientos Informados")

st.markdown("""
Sube el archivo **Excel** con los datos de los centros y el **Word modelo** con los placeholders.
Se generará un documento Word por cada investigador y se descargará todo en un ZIP.
""")

excel_file = st.file_uploader("📊 Subir Excel con datos", type=["xlsx"])
word_template = st.file_uploader("📄 Subir documento Word modelo", type=["docx"])

if excel_file and word_template:
    if st.button("⚙️ Generar documentos"):
        with st.spinner("Generando documentos..."):
            zip_file = generar_documentos(excel_file, word_template)

        st.success("✅ Documentos generados correctamente")
        with open(zip_file, "rb") as f:
            st.download_button(
                label="⬇️ Descargar ZIP",
                data=f,
                file_name="documentos_consentimiento.zip",
                mime="application/zip"
            )


